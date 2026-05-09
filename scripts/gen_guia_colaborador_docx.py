"""Genera Guia_Colaborador_TPT.docx con estilos profesionales (espanol neutro)."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Colores MSO
MORADO = RGBColor(0x3D, 0x0C, 0x4B)
MORADO_CLARO = RGBColor(0x6B, 0x1D, 0x7B)
NARANJO = RGBColor(0xF5, 0x82, 0x20)
TEXTO = RGBColor(0x2D, 0x37, 0x48)
TEXTO_SEC = RGBColor(0x71, 0x80, 0x96)

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = TEXTO

# ==========================
# HELPERS
# ==========================
def add_heading(text, level=1, color=None):
    if color is None:
        color = MORADO
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    if level == 1:
        run.font.size = Pt(22)
    elif level == 2:
        run.font.size = Pt(17)
    else:
        run.font.size = Pt(14)
    run.bold = True
    return h

def add_para(text, size=11, bold=False, color=None, align=None):
    if color is None:
        color = TEXTO
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    return p

def add_bullet(text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.font.color.rgb = TEXTO
    return p

def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def add_callout(tipo, texto):
    colores = {
        'tip':     ('E3F2FD', '1976D2', 'Consejo'),
        'warning': ('FFF3E0', 'F57C00', 'Importante'),
        'success': ('E8F5E9', '388E3C', 'Tip')
    }
    bg, border_hex, titulo = colores[tipo]
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg)
    cell.paragraphs[0].clear()
    p_tit = cell.paragraphs[0]
    r_tit = p_tit.add_run(titulo + "\n")
    r_tit.bold = True
    r_tit.font.color.rgb = RGBColor.from_string(border_hex)
    r_tit.font.size = Pt(11)
    p_body = cell.add_paragraph()
    r_body = p_body.add_run(texto)
    r_body.font.size = Pt(10.5)
    r_body.font.color.rgb = TEXTO
    doc.add_paragraph()

def add_step(numero, texto):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(1.2)
    table.columns[1].width = Cm(15.3)
    cell_num = table.cell(0, 0)
    cell_txt = table.cell(0, 1)
    set_cell_bg(cell_num, '3D0C4B')
    cell_num.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_num.paragraphs[0].clear()
    p_num = cell_num.paragraphs[0]
    p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_num = p_num.add_run(str(numero))
    r_num.bold = True
    r_num.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r_num.font.size = Pt(14)

    cell_txt.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_txt.paragraphs[0].clear()
    p_txt = cell_txt.paragraphs[0]
    r_txt = p_txt.add_run(texto)
    r_txt.font.size = Pt(11)
    r_txt.font.color.rgb = TEXTO
    doc.add_paragraph()

def add_table_header(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, '3D0C4B')
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(11)
    for row_idx, row_data in enumerate(rows, 1):
        row = table.rows[row_idx]
        for col_idx, val in enumerate(row_data):
            cell = row.cells[col_idx]
            if row_idx % 2 == 0:
                set_cell_bg(cell, 'FAFAFA')
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10.5)
            r.font.color.rgb = TEXTO
    doc.add_paragraph()

# ==========================
# PORTADA
# ==========================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MSO CHILE')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = MORADO

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Guía del Colaborador')
run.bold = True
run.font.size = Pt(40)
run.font.color.rgb = MORADO

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Plataforma de Transferencia al Puesto de Trabajo (TPT)')
run.font.size = Pt(16)
run.font.color.rgb = NARANJO
run.italic = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'Esta guía te acompaña durante el programa. Vas a aprender cómo ingresar a '
    'la plataforma, responder la coevaluación de tu líder, ver tu progreso, '
    'consultar los recursos disponibles y recibir el feedback que tu líder te '
    'entrega. El objetivo es que tengas todo lo necesario a la mano.'
)
run.font.size = Pt(12)
run.font.color.rgb = TEXTO_SEC
run.italic = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Versión 1.0  ·  Abril 2026')
run.font.size = Pt(11)
run.font.color.rgb = TEXTO_SEC

doc.add_page_break()

# ==========================
# CONTENIDO
# ==========================
add_heading('Contenido', level=1)
toc_items = [
    '1. ¿Qué es la plataforma TPT?',
    '2. Tu rol como colaborador',
    '3. Primer acceso a la plataforma',
    '4. ¿Olvidaste tu contraseña?',
    '5. Tu pantalla principal: Mi Programa',
    '6. Cómo responder la coevaluación de tu líder',
    '7. Escala de niveles (1 a 4)',
    '8. Mi Progreso: seguir tu avance',
    '9. Feedback recibido de tu líder',
    '10. Archivos y recursos del programa',
    '11. Reportar incidencias',
    '12. Consejos para aprovechar el programa',
    '13. Preguntas frecuentes',
    '14. Soporte y contacto',
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(item)
    run.font.size = Pt(12)
    run.font.color.rgb = MORADO

doc.add_page_break()

# ==========================
# 1. ¿QUE ES?
# ==========================
add_heading('1. ¿Qué es la plataforma TPT?', level=1)
add_para(
    'La plataforma de Transferencia al Puesto de Trabajo (TPT) es la herramienta '
    'que MSO usa para acompañar programas de desarrollo de competencias en '
    'equipos de trabajo. A través de ella, tú y tu líder responden evaluaciones, '
    'reciben feedback y acceden a los materiales del programa.'
)

add_heading('¿Para qué sirve?', level=2, color=MORADO_CLARO)
add_bullet('Medir tus percepciones sobre las competencias y conductas de tu líder.')
add_bullet('Registrar el feedback que tu líder te entrega a lo largo del programa.')
add_bullet('Consultar los recursos y lecturas del programa en un solo lugar.')
add_bullet('Ver tu progreso en las actividades planificadas.')

# ==========================
# 2. ROL COLABORADOR
# ==========================
add_heading('2. Tu rol como colaborador', level=1)
add_para(
    'Como colaborador, participas acompañando a tu líder en el programa. Tu '
    'rol principal es dar una mirada externa y honesta sobre las competencias '
    'y conductas de tu líder, a través de dos coevaluaciones.'
)

add_heading('¿Qué se espera de ti?', level=2, color=MORADO_CLARO)
add_bullet('Responder la coevaluación inicial al comienzo del programa.')
add_bullet('Responder la coevaluación final al cierre del programa.')
add_bullet('Leer y considerar el feedback que tu líder te entrega.')
add_bullet('Consultar los recursos cuando lo necesites.')

add_callout('success',
    'Tu participación es fundamental. Las coevaluaciones permiten identificar '
    'brechas entre cómo se ve tu líder y cómo lo ven las personas de su equipo. '
    'Sin tu aporte, la medición queda incompleta.')

# ==========================
# 3. PRIMER ACCESO
# ==========================
add_heading('3. Primer acceso a la plataforma', level=1)
add_para(
    'Cuando el administrador del programa carga tus datos, recibes un correo '
    'de bienvenida con tus credenciales. Si aún no lo recibes, revisa tu carpeta '
    'de spam o correo no deseado.'
)

add_heading('Pasos para ingresar', level=2, color=MORADO_CLARO)
add_step(1, 'Abre un navegador (Chrome, Edge, Firefox o Safari) e ingresa a: https://plataforma.msochile.cl')
add_step(2, 'Escribe tu correo electrónico en el campo correspondiente.')
add_step(3, 'Escribe tu contraseña. Si es tu primer acceso, usa la que recibiste por correo.')
add_step(4, 'Haz click en "Ingresar".')

add_callout('tip',
    'Te recomendamos cambiar tu contraseña en el primer acceso. Puedes hacerlo '
    'desde la opción "¿Olvidaste tu contraseña?" del login y crear una nueva que '
    'sea fácil de recordar para ti.')

# ==========================
# 4. RECUPERAR
# ==========================
add_heading('4. ¿Olvidaste tu contraseña?', level=1)
add_para('Si no recuerdas tu contraseña, sigue estos pasos:')

add_step(1, 'En la pantalla de login, haz click en "¿Olvidaste tu contraseña?"')
add_step(2, 'Escribe tu correo electrónico y haz click en "Enviar instrucciones".')
add_step(3, 'Recibirás un correo desde "MSO Plataforma <no_reply@msochile.cl>" con un link para restablecer tu contraseña.')
add_step(4, 'El link te lleva a una pantalla donde puedes definir tu nueva contraseña. Elige una combinación segura de al menos 6 caracteres.')

add_callout('warning',
    'El link de recuperación es válido por una hora y solo puede usarse una vez '
    'por razones de seguridad. Si expira, solicita otro.')

# ==========================
# 5. MI PROGRAMA
# ==========================
add_heading('5. Tu pantalla principal: Mi Programa', level=1)
add_para(
    'Al iniciar sesión, tu pantalla principal es Mi Programa. Allí ves toda la '
    'información relevante del programa en el que participas.'
)

add_heading('¿Qué encuentras en Mi Programa?', level=2, color=MORADO_CLARO)
add_bullet('Datos generales del programa: nombre, cliente, fechas de inicio y término.')
add_bullet('Nombre del líder al que estás asociado.')
add_bullet('Las fases del programa con sus pasos (medición inicial, medición final, comparativo).')
add_bullet('Estado de cada paso: completada, en curso, pendiente o bloqueada.')
add_bullet('Botón "Responder" cuando una coevaluación esté disponible.')

add_callout('tip',
    'Esta vista es tu punto de partida. Si un paso aparece como "Pendiente", '
    'significa que debes completarlo. Cuando dice "Completada", ya está cerrado.')

# ==========================
# 6. COEVALUACION
# ==========================
add_heading('6. Cómo responder la coevaluación de tu líder', level=1)
add_para(
    'La coevaluación es el instrumento que usas para evaluar las competencias '
    'y conductas de tu líder. Vas a responder dos veces: al comienzo del '
    'programa (medición inicial) y al cierre (medición final).'
)

add_heading('Pasos para responder', level=2, color=MORADO_CLARO)
add_step(1, 'En el menú lateral, haz click en "Mis Evaluaciones".')
add_step(2, 'Verás dos secciones: "Encuestas Pendientes" y "Encuestas Completadas".')
add_step(3, 'En la encuesta pendiente, haz click en "Responder".')
add_step(4, 'Lee las instrucciones con atención. Cada pregunta se refiere a una conducta observable de tu líder.')
add_step(5, 'Responde cada pregunta seleccionando el nivel que mejor describe lo que has observado (escala de 1 a 4).')
add_step(6, 'Usa la barra de progreso para saber cuántas preguntas llevas respondidas.')
add_step(7, 'Al terminar, haz click en "Enviar Respuestas". Recibirás un correo de confirmación.')

add_heading('¿Te equivocaste al responder?', level=2, color=MORADO_CLARO)
add_callout('success',
    'Si ya enviaste una coevaluación y necesitas corregirla, ve a Mis '
    'Evaluaciones → Encuestas Completadas y haz click en "Rehacer evaluación". '
    'Tus respuestas anteriores serán reemplazadas por las nuevas.')

add_callout('warning',
    'Responde con honestidad y basándote en situaciones que hayas observado. El '
    'objetivo no es calificar de buena o mala manera a tu líder, sino ayudarlo '
    'a identificar oportunidades concretas de desarrollo.')

# ==========================
# 7. ESCALA
# ==========================
add_heading('7. Escala de niveles (1 a 4)', level=1)
add_para(
    'La mayoría de las preguntas de la coevaluación usan una escala de 4 niveles. '
    'Esta escala te ayuda a precisar qué tanto observas que tu líder aplica cada '
    'conducta.'
)

add_table_header(
    ['Nivel', 'Descripción'],
    [
        ['1', 'Conoce el concepto pero no lo aplica.'],
        ['2', 'Aplica con orientación o apoyo.'],
        ['3', 'Aplica de forma consistente y autónoma.'],
        ['4', 'Es referente y modelo para otros.'],
    ]
)

add_callout('tip',
    'Piensa en situaciones concretas del último tiempo antes de escoger un nivel. '
    'Si no tienes evidencia de una conducta, prefiere un nivel más bajo antes que '
    'asumir que se cumple.')

# ==========================
# 8. MI PROGRESO
# ==========================
add_heading('8. Mi Progreso: seguir tu avance', level=1)
add_para(
    'En la sección "Mi Progreso" ves el resumen de tu participación en el '
    'programa: encuestas respondidas, actividades completadas y pasos pendientes.'
)

add_heading('¿Qué muestra esta sección?', level=2, color=MORADO_CLARO)
add_bullet('Cuántas encuestas has respondido y cuántas faltan.')
add_bullet('Fecha de cada respuesta enviada.')
add_bullet('Estado general de tu participación.')

# ==========================
# 9. FEEDBACK
# ==========================
add_heading('9. Feedback recibido de tu líder', level=1)
add_para(
    'Durante el programa, tu líder puede registrar feedback sobre ti en tres '
    'dimensiones: fortalezas, aspectos a reforzar y recomendaciones. Cada vez '
    'que lo hace, queda disponible para que lo consultes.'
)

add_heading('Consultar tu feedback', level=2, color=MORADO_CLARO)
add_step(1, 'En el menú lateral, haz click en "Feedback Recibido".')
add_step(2, 'Verás la lista de feedback recibidos ordenados por fecha.')
add_step(3, 'Al abrir cada uno, encuentras las tres dimensiones: fortaleza observada, aspecto a reforzar y recomendación.')

add_callout('tip',
    'Vuelve a esta sección cada cierto tiempo. El feedback es más útil cuando '
    'lo revisas con calma y lo llevas a acciones concretas en tu día a día.')

# ==========================
# 10. RECURSOS
# ==========================
add_heading('10. Archivos y recursos del programa', level=1)
add_para(
    'El administrador del programa puede subir documentos, lecturas y '
    'materiales complementarios. Allí puedes encontrar guías, artículos y '
    'cualquier recurso útil para tu desarrollo.'
)
add_step(1, 'Ve a "Archivos y Recursos" en el menú lateral.')
add_step(2, 'Verás la lista de archivos disponibles. Haz click en cada uno para descargarlo o visualizarlo.')

# ==========================
# 11. INCIDENCIAS
# ==========================
add_heading('11. Reportar incidencias', level=1)
add_para(
    'Si encuentras un error técnico, tienes una duda o quieres sugerir una '
    'mejora, puedes reportarlo directamente desde la plataforma.'
)
add_step(1, 'En el menú lateral, haz click en "Incidencias".')
add_step(2, 'Completa los campos: categoría, título, descripción del problema y, si corresponde, adjunta una captura de pantalla.')
add_step(3, 'Haz click en "Reportar". El equipo de soporte recibirá tu reporte y te contactará si es necesario.')

# ==========================
# 12. CONSEJOS
# ==========================
add_heading('12. Consejos para aprovechar el programa', level=1)
add_bullet('Reserva un momento tranquilo para responder cada coevaluación. No la dejes para último minuto.')
add_bullet('Responde con honestidad, basándote en situaciones concretas que hayas observado.')
add_bullet('Cuando recibas feedback, léelo con apertura y pregúntate qué acciones puedes tomar.')
add_bullet('Revisa los recursos del programa de a poco, no hace falta leer todo de una sola vez.')
add_bullet('Si algo no te queda claro, consulta al administrador o reporta una incidencia.')

# ==========================
# 13. FAQ
# ==========================
add_heading('13. Preguntas frecuentes', level=1)

faqs = [
    ('No me llegó el correo de bienvenida, ¿qué hago?',
     'Revisa tu carpeta de spam o correo no deseado. Si tampoco está allí, contacta al administrador del programa para que verifique el envío.'),
    ('El link de reset de contraseña dice que expiró',
     'Los links de reset son válidos por una hora y solo pueden usarse una vez. Si expira, solicita otro desde "¿Olvidaste tu contraseña?" en el login.'),
    ('No veo mi coevaluación disponible',
     'Las encuestas se activan según el cronograma del programa. Si la fecha ya pasó y no la ves, contacta al administrador del programa para que revise el estado de activación.'),
    ('Me equivoqué al responder una pregunta, ¿puedo corregir?',
     'Sí. Ve a Mis Evaluaciones → Encuestas Completadas y haz click en "Rehacer evaluación". Tus respuestas anteriores serán reemplazadas por las nuevas.'),
    ('¿Puedo responder desde el celular?',
     'Sí, la plataforma funciona en celulares y tablets. Para coevaluaciones largas recomendamos computador o tablet para mayor comodidad.'),
    ('¿Mis respuestas son confidenciales?',
     'Las coevaluaciones se usan en los informes del programa. El análisis es agregado: tu líder recibe un informe con brechas de percepción, pero no identifica individualmente a cada colaborador.'),
    ('No aparece mi líder asignado',
     'Contacta al administrador del programa para que verifique la asignación. Es importante que quede correctamente registrado para que tu coevaluación sume al informe correcto.'),
    ('¿Puedo ver las respuestas de mi líder sobre sí mismo?',
     'No. La autoevaluación de tu líder es confidencial. Lo que sí verás es el feedback que tu líder te entregue formalmente a través de la plataforma.'),
]
for q, a in faqs:
    p = doc.add_paragraph()
    r = p.add_run(q)
    r.bold = True
    r.font.color.rgb = MORADO
    r.font.size = Pt(12)
    add_para(a, size=11)

# ==========================
# 14. SOPORTE
# ==========================
add_heading('14. Soporte y contacto', level=1)
add_para('Si necesitas ayuda o tienes una duda que no aparece en esta guía, cuentas con estos canales:')
add_table_header(
    ['Canal', 'Detalle'],
    [
        ['Administrador del programa', 'Consultas sobre fechas, activación de encuestas o asignación de líder.'],
        ['Soporte técnico MSO', 'Problemas de acceso, errores técnicos o sugerencias sobre la plataforma.'],
        ['Módulo de Incidencias', 'Reporta dentro de la plataforma cualquier error que encuentres.'],
    ]
)

# ==========================
# CIERRE
# ==========================
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('MSO Chile · Modelos y Soluciones Organizacionales')
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = MORADO

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Plataforma de Transferencia al Puesto de Trabajo (TPT)')
r.font.size = Pt(10)
r.font.color.rgb = TEXTO_SEC

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('https://plataforma.msochile.cl  ·  Abril 2026')
r.font.size = Pt(10)
r.font.color.rgb = TEXTO_SEC

out = r"c:\Users\the_r\Documents\GitHub\plataforma-mso\docs\Guia_Colaborador_TPT.docx"
doc.save(out)
print("OK:", out)
import os
print("Tamano:", round(os.path.getsize(out)/1024, 1), "KB")
