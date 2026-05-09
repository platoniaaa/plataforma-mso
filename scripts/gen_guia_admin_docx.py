"""Genera Guia_Administrador_TPT.docx con estilos profesionales (espanol neutro)."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Colores MSO
MORADO = RGBColor(0x3D, 0x0C, 0x4B)
MORADO_CLARO = RGBColor(0x6B, 0x1D, 0x7B)
NARANJO = RGBColor(0xF5, 0x82, 0x20)
TEXTO = RGBColor(0x2D, 0x37, 0x48)
TEXTO_SEC = RGBColor(0x71, 0x80, 0x96)

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = TEXTO

# ==========================
# HELPERS
# ==========================
def add_heading(text, level=1, color=None):
    if color is None:
        color = MORADO
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    if level == 1:
        run.font.size = Pt(22)
    elif level == 2:
        run.font.size = Pt(17)
    else:
        run.font.size = Pt(14)
    run.bold = True
    return h

def add_para(text, size=11, bold=False, color=None, align=None):
    if color is None:
        color = TEXTO
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    return p

def add_bullet(text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.font.color.rgb = TEXTO
    return p

def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def add_callout(tipo, texto):
    colores = {
        'tip':     ('E3F2FD', '1976D2', 'Consejo'),
        'warning': ('FFF3E0', 'F57C00', 'Importante'),
        'success': ('E8F5E9', '388E3C', 'Tip')
    }
    bg, border_hex, titulo = colores[tipo]
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg)
    cell.paragraphs[0].clear()
    p_tit = cell.paragraphs[0]
    r_tit = p_tit.add_run(titulo + "\n")
    r_tit.bold = True
    r_tit.font.color.rgb = RGBColor.from_string(border_hex)
    r_tit.font.size = Pt(11)
    p_body = cell.add_paragraph()
    r_body = p_body.add_run(texto)
    r_body.font.size = Pt(10.5)
    r_body.font.color.rgb = TEXTO
    doc.add_paragraph()

def add_step(numero, texto):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(1.2)
    table.columns[1].width = Cm(15.3)
    cell_num = table.cell(0, 0)
    cell_txt = table.cell(0, 1)
    set_cell_bg(cell_num, '3D0C4B')
    cell_num.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_num.paragraphs[0].clear()
    p_num = cell_num.paragraphs[0]
    p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_num = p_num.add_run(str(numero))
    r_num.bold = True
    r_num.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r_num.font.size = Pt(14)

    cell_txt.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_txt.paragraphs[0].clear()
    p_txt = cell_txt.paragraphs[0]
    r_txt = p_txt.add_run(texto)
    r_txt.font.size = Pt(11)
    r_txt.font.color.rgb = TEXTO
    doc.add_paragraph()

def add_table_header(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, '3D0C4B')
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(11)
    for row_idx, row_data in enumerate(rows, 1):
        row = table.rows[row_idx]
        for col_idx, val in enumerate(row_data):
            cell = row.cells[col_idx]
            if row_idx % 2 == 0:
                set_cell_bg(cell, 'FAFAFA')
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10.5)
            r.font.color.rgb = TEXTO
    doc.add_paragraph()

# ==========================
# PORTADA
# ==========================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MSO CHILE')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = MORADO

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Guía del Administrador')
run.bold = True
run.font.size = Pt(40)
run.font.color.rgb = MORADO

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Plataforma de Transferencia al Puesto de Trabajo (TPT)')
run.font.size = Pt(16)
run.font.color.rgb = NARANJO
run.italic = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'Esta guía te acompaña en la gestión integral de la plataforma: desde la '
    'creación de clientes y programas, hasta la carga de participantes, '
    'activación de encuestas, seguimiento del avance y generación de informes. '
    'Está pensada para el equipo de MSO que administra el ciclo completo de '
    'cada programa.'
)
run.font.size = Pt(12)
run.font.color.rgb = TEXTO_SEC
run.italic = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Versión 1.0  ·  Abril 2026')
run.font.size = Pt(11)
run.font.color.rgb = TEXTO_SEC

doc.add_page_break()

# ==========================
# CONTENIDO
# ==========================
add_heading('Contenido', level=1)
toc_items = [
    '1. Acceso y rol de administrador',
    '2. Panel principal: Dashboard',
    '3. Gestión de clientes',
    '4. Gestión de programas',
    '5. Carga de competencias y conductas',
    '6. Carta Gantt del programa',
    '7. Carga de participantes (líderes y colaboradores)',
    '8. Editor de encuestas y activación',
    '9. Seguimiento de respuestas',
    '10. Feedback entre líder y colaborador',
    '11. Recursos y archivos del programa',
    '12. Informes con inteligencia artificial',
    '13. Gestión de correos enviados',
    '14. Atención de incidencias reportadas',
    '15. Notificaciones y recordatorios',
    '16. Buenas prácticas y recomendaciones',
    '17. Preguntas frecuentes',
    '18. Soporte y contacto',
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(item)
    run.font.size = Pt(12)
    run.font.color.rgb = MORADO

doc.add_page_break()

# ==========================
# 1. ACCESO
# ==========================
add_heading('1. Acceso y rol de administrador', level=1)
add_para(
    'El perfil de administrador tiene permisos amplios sobre la plataforma: '
    'crear clientes, crear programas, cargar participantes, activar encuestas, '
    'generar informes y gestionar incidencias. Es el rol que ejecuta el equipo '
    'de MSO durante todo el ciclo de un programa.'
)

add_heading('Pasos para ingresar', level=2, color=MORADO_CLARO)
add_step(1, 'Abre un navegador (Chrome, Edge, Firefox o Safari) e ingresa a: https://plataforma.msochile.cl')
add_step(2, 'Ingresa el correo electrónico corporativo con el que se creó tu cuenta de administrador.')
add_step(3, 'Ingresa tu contraseña y haz click en "Ingresar".')
add_step(4, 'Si olvidaste tu contraseña, utiliza la opción "¿Olvidaste tu contraseña?" del login. Recibirás un correo con un link válido por una hora.')

add_callout('warning',
    'Solo el equipo de MSO debe tener cuentas de administrador. No compartas tu '
    'contraseña ni crees cuentas adicionales sin autorización.')

# ==========================
# 2. DASHBOARD
# ==========================
add_heading('2. Panel principal: Dashboard', level=1)
add_para(
    'Al iniciar sesión, la primera pantalla es el Dashboard. Desde allí tienes '
    'una vista general de la actividad en la plataforma.'
)

add_heading('¿Qué muestra el Dashboard?', level=2, color=MORADO_CLARO)
add_bullet('Totales de clientes, programas activos y participantes.')
add_bullet('Listado de programas en curso con su avance.')
add_bullet('Accesos rápidos a las secciones más usadas.')
add_bullet('Indicadores de encuestas pendientes o vencidas.')

add_callout('tip',
    'Usa el Dashboard como punto de partida diario para detectar programas que '
    'requieren atención inmediata, por ejemplo, fechas de encuesta próximas a vencer.')

# ==========================
# 3. CLIENTES
# ==========================
add_heading('3. Gestión de clientes', level=1)
add_para(
    'Cada programa se asocia a un cliente. Los clientes son las organizaciones '
    'a las que MSO presta servicio (por ejemplo, Sodexo).'
)

add_heading('Crear un cliente', level=2, color=MORADO_CLARO)
add_step(1, 'En el menú lateral, haz click en "Clientes".')
add_step(2, 'Haz click en el botón "Nuevo Cliente".')
add_step(3, 'Completa los datos: nombre, razón social, RUT, industria y contacto principal.')
add_step(4, 'Guarda los cambios. El cliente queda disponible para asociar nuevos programas.')

add_heading('Editar o desactivar un cliente', level=2, color=MORADO_CLARO)
add_bullet('Para editar, haz click sobre la tarjeta del cliente y modifica los campos necesarios.')
add_bullet('Para desactivar, usa la opción correspondiente. Los programas históricos del cliente se mantienen disponibles para consulta.')

add_callout('warning',
    'No elimines clientes con programas en curso. Si necesitas dejar de atenderlo, '
    'primero cierra sus programas activos.')

# ==========================
# 4. PROGRAMAS
# ==========================
add_heading('4. Gestión de programas', level=1)
add_para(
    'Un programa es una instancia concreta de intervención para un cliente. '
    'Tiene fechas de inicio y término, competencias, participantes y encuestas '
    'asociadas.'
)

add_heading('Crear un programa', level=2, color=MORADO_CLARO)
add_step(1, 'En el menú lateral, haz click en "Programas".')
add_step(2, 'Haz click en "Nuevo Programa".')
add_step(3, 'Selecciona el cliente al que pertenece.')
add_step(4, 'Completa: nombre del programa, objetivo, fechas de inicio y término.')
add_step(5, 'Guarda. El programa se crea vacío y listo para recibir competencias, cronograma y participantes.')

add_heading('Panel del programa', level=2, color=MORADO_CLARO)
add_para(
    'Al abrir un programa existente, accedes a su panel. Allí encuentras pestañas '
    'para administrar todas sus secciones:'
)
add_table_header(
    ['Pestaña', 'Descripción'],
    [
        ['Competencias', 'Competencias y conductas evaluadas en el programa.'],
        ['Cronograma / Gantt', 'Actividades y fechas del programa en formato Carta Gantt.'],
        ['Participantes', 'Líderes y colaboradores que participan del programa.'],
        ['Encuestas', 'Encuestas creadas, su estado y avance.'],
        ['Editor de Encuesta', 'Configuración fina de una encuesta puntual.'],
        ['Informes', 'Informes generados por la plataforma con IA.'],
    ]
)

# ==========================
# 5. COMPETENCIAS
# ==========================
add_heading('5. Carga de competencias y conductas', level=1)
add_para(
    'Cada programa tiene un conjunto de competencias que a su vez agrupan '
    'conductas observables. Estas son la base de las evaluaciones.'
)

add_heading('Carga masiva desde Excel', level=2, color=MORADO_CLARO)
add_step(1, 'Abre la pestaña "Competencias" del programa.')
add_step(2, 'Descarga la plantilla Excel haciendo click en "Descargar plantilla".')
add_step(3, 'Completa la plantilla: una fila por conducta, indicando la competencia a la que pertenece.')
add_step(4, 'Guarda el archivo y súbelo con el botón "Cargar Excel".')
add_step(5, 'La plataforma valida y crea las competencias y conductas automáticamente.')

add_heading('Carga manual', level=2, color=MORADO_CLARO)
add_bullet('Desde la misma pestaña, puedes crear competencias y conductas de forma individual.')
add_bullet('Útil cuando necesitas ajustar un detalle específico sin volver a cargar todo el Excel.')

add_callout('tip',
    'Antes de subir el Excel, valida que no existan nombres duplicados. La plantilla '
    'oficial incluye una hoja con instrucciones para evitar errores comunes.')

# ==========================
# 6. GANTT
# ==========================
add_heading('6. Carta Gantt del programa', level=1)
add_para(
    'La Carta Gantt define las actividades del programa, sus fechas y su orden. '
    'Es la línea de tiempo que todos los participantes siguen.'
)

add_heading('Cargar la Carta Gantt', level=2, color=MORADO_CLARO)
add_step(1, 'Abre la pestaña "Cronograma" o "Gantt" del programa.')
add_step(2, 'Descarga la plantilla Excel.')
add_step(3, 'Completa cada actividad con: nombre, fecha de inicio, fecha de término y tipo.')
add_step(4, 'Guarda y sube el archivo desde "Cargar Gantt".')

add_callout('warning',
    'La Carta Gantt puede actualizar automáticamente la fecha de cierre de las '
    'encuestas asociadas. Revisa el editor de encuesta después de cada carga para '
    'confirmar que las fechas están como esperas.')

# ==========================
# 7. PARTICIPANTES
# ==========================
add_heading('7. Carga de participantes (líderes y colaboradores)', level=1)
add_para(
    'Los participantes del programa son los líderes y sus colaboradores asignados. '
    'Cada uno recibe un correo de bienvenida con sus credenciales cuando se cargan.'
)

add_heading('Carga masiva por Excel', level=2, color=MORADO_CLARO)
add_step(1, 'Abre la pestaña "Participantes" del programa.')
add_step(2, 'Descarga la plantilla de carga (encuentras un enlace en la misma pestaña).')
add_step(3, 'Completa una fila por persona: nombre, correo, rol (líder o colaborador), cargo y líder asociado cuando corresponda.')
add_step(4, 'Sube el archivo. La plataforma crea los usuarios, genera contraseñas aleatorias y envía el correo de bienvenida.')

add_heading('Tabla de participantes', level=2, color=MORADO_CLARO)
add_para('La tabla de participantes ofrece las siguientes herramientas:')
add_bullet('Ícono de ojo para visualizar la contraseña asignada a cada usuario.')
add_bullet('Botón para eliminar participantes de forma individual.')
add_bullet('Botón "Eliminar todos" para vaciar el programa cuando necesitas rehacer la carga.')
add_bullet('Columna de contraseña visible tanto para líderes como para colaboradores.')

add_callout('warning',
    'El botón "Eliminar todos" borra participantes del programa y sus cuentas '
    'asociadas. Úsalo solamente cuando vas a volver a cargar el Excel desde cero.')

# ==========================
# 8. EDITOR DE ENCUESTAS
# ==========================
add_heading('8. Editor de encuestas y activación', level=1)
add_para(
    'Cada programa tiene varias encuestas: autoevaluaciones y coevaluaciones, '
    'inicial y final. Desde el editor puedes controlar cuándo están disponibles '
    'para los participantes.'
)

add_heading('Abrir el editor de una encuesta', level=2, color=MORADO_CLARO)
add_step(1, 'Abre la pestaña "Encuestas" del programa.')
add_step(2, 'Haz click sobre la encuesta que quieres configurar.')
add_step(3, 'Se abre el editor con su estado actual, fecha de cierre y preguntas.')

add_heading('Activar y cerrar una encuesta', level=2, color=MORADO_CLARO)
add_bullet('Usa el botón "Guardar cambios" luego de editar fecha u otros campos.')
add_bullet('Usa "Reabrir encuesta" para volver a dejar disponible una encuesta ya cerrada.')
add_bullet('El campo "Fecha de cierre" se puede ingresar manualmente y también se sincroniza con la Carta Gantt.')

add_callout('tip',
    'Si necesitas que un participante corrija una respuesta enviada, puedes '
    'pedirle que use la opción "Rehacer evaluación" en su perfil. No hace falta '
    'reabrir la encuesta completa.')

# ==========================
# 9. SEGUIMIENTO
# ==========================
add_heading('9. Seguimiento de respuestas', level=1)
add_para(
    'Desde la pestaña "Encuestas" del programa puedes monitorear cuántos '
    'participantes han respondido cada encuesta y cuántos están pendientes.'
)
add_bullet('Barra de progreso por encuesta con el porcentaje de avance.')
add_bullet('Listado de participantes pendientes para seguimiento manual.')
add_bullet('Posibilidad de enviar un recordatorio puntual.')

add_callout('tip',
    'La plataforma también envía recordatorios automáticos a los pendientes cada '
    'día a las 12:00 UTC. No necesitas gatillar el envío manualmente salvo casos '
    'excepcionales.')

# ==========================
# 10. FEEDBACK
# ==========================
add_heading('10. Feedback entre líder y colaborador', level=1)
add_para(
    'El módulo de feedback permite que cada líder registre observaciones '
    'estructuradas sobre su colaborador. Como administrador, puedes consultar '
    'el histórico para asegurar su uso.'
)

add_heading('¿Qué registra el líder?', level=2, color=MORADO_CLARO)
add_bullet('Fortaleza observada.')
add_bullet('Aspecto a reforzar.')
add_bullet('Recomendación concreta.')

add_callout('success',
    'El feedback queda guardado en la plataforma y el colaborador puede '
    'consultarlo desde su propia sesión en "Feedback Recibido".')

# ==========================
# 11. RECURSOS
# ==========================
add_heading('11. Recursos y archivos del programa', level=1)
add_para(
    'Puedes subir documentos, lecturas, presentaciones u otros recursos para los '
    'participantes del programa.'
)
add_step(1, 'Abre la sección de "Archivos y Recursos" del programa.')
add_step(2, 'Haz click en "Subir archivo" y selecciona el documento desde tu computador.')
add_step(3, 'Agrega un nombre descriptivo y, si corresponde, una breve descripción.')
add_step(4, 'Guarda. El recurso queda disponible para todos los participantes del programa.')

# ==========================
# 12. INFORMES IA
# ==========================
add_heading('12. Informes con inteligencia artificial', level=1)
add_para(
    'La plataforma utiliza un asistente de IA para generar informes a partir de '
    'las respuestas de las encuestas. Todos los informes pasan por revisión '
    'antes de ser entregados al cliente.'
)

add_heading('Tipos de informe disponibles', level=2, color=MORADO_CLARO)
add_table_header(
    ['Informe', 'Descripción'],
    [
        ['Individual Inicial',   'Resumen personal de cada líder al comenzar el programa.'],
        ['Consolidado Inicial',  'Vista agregada del grupo de líderes en la medición inicial.'],
        ['Individual Final',     'Resumen personal al cierre del programa.'],
        ['Consolidado Final',    'Vista agregada del grupo al cierre del programa.'],
        ['Brechas / Comparativo','Comparativo entre la medición inicial y la final.'],
    ]
)

add_heading('Generar un informe', level=2, color=MORADO_CLARO)
add_step(1, 'Abre la pestaña "Informes" del programa.')
add_step(2, 'Selecciona el tipo de informe y el participante (o el grupo completo si es consolidado).')
add_step(3, 'Haz click en "Generar". El asistente procesa las respuestas y devuelve un borrador.')
add_step(4, 'Revisa el contenido, ajusta si corresponde y descarga en el formato final.')

add_callout('warning',
    'Siempre revisa el informe antes de entregarlo al cliente. La IA es un apoyo '
    'para acelerar la redacción, no un reemplazo del juicio profesional.')

# ==========================
# 13. CORREOS
# ==========================
add_heading('13. Gestión de correos enviados', level=1)
add_para(
    'Toda comunicación que envía la plataforma queda registrada en la sección '
    '"Correos". Desde allí puedes auditar qué se envió, a quién y con qué estado.'
)

add_heading('¿Qué correos se registran?', level=2, color=MORADO_CLARO)
add_bullet('Correo de bienvenida con credenciales.')
add_bullet('Encuesta disponible para responder.')
add_bullet('Recordatorios automáticos de encuestas pendientes.')
add_bullet('Confirmación de respuesta enviada.')
add_bullet('Notificación al líder cuando su colaborador lo coevaluó.')
add_bullet('Solicitud y confirmación de restablecimiento de contraseña.')

add_heading('Estados posibles', level=2, color=MORADO_CLARO)
add_table_header(
    ['Estado', 'Significado'],
    [
        ['enviado', 'Todos los destinatarios recibieron el correo correctamente.'],
        ['parcial', 'Algunos destinatarios recibieron, otros fallaron. Revisa la columna de error.'],
        ['fallido', 'Ningún destinatario recibió el correo. Revisa la columna de error.'],
    ]
)

add_callout('tip',
    'Si un usuario reclama que no recibió un correo, busca su caso en esta sección. '
    'Los IDs de Resend te permiten validar el envío directamente en el dashboard '
    'del proveedor de correo.')

# ==========================
# 14. INCIDENCIAS
# ==========================
add_heading('14. Atención de incidencias reportadas', level=1)
add_para(
    'Los usuarios pueden reportar problemas o dudas desde la plataforma. Como '
    'administrador puedes revisarlas, asignarlas y marcarlas como resueltas.'
)

add_heading('Flujo sugerido', level=2, color=MORADO_CLARO)
add_step(1, 'Abre la sección "Incidencias" en el menú lateral.')
add_step(2, 'Revisa los reportes abiertos. Cada uno incluye categoría, título, descripción y, si hubo, captura adjunta.')
add_step(3, 'Responde al usuario, resuelve el problema o escala al equipo técnico según corresponda.')
add_step(4, 'Marca la incidencia como resuelta una vez cerrada.')

# ==========================
# 15. NOTIFICACIONES
# ==========================
add_heading('15. Notificaciones y recordatorios', level=1)
add_para(
    'El sistema envía recordatorios automáticos todos los días a las 12:00 UTC '
    'a los participantes con encuestas pendientes. Además, cada acción clave '
    'genera una notificación interna.'
)

add_heading('Buenas prácticas para no saturar usuarios', level=2, color=MORADO_CLARO)
add_bullet('Evita activar varias encuestas el mismo día si no es necesario.')
add_bullet('Coordina con el cliente la ventana de respuesta (mínimo una semana).')
add_bullet('Si ves muchos pendientes, prefiere recordatorios puntuales antes que reenvíos masivos.')

# ==========================
# 16. BUENAS PRACTICAS
# ==========================
add_heading('16. Buenas prácticas y recomendaciones', level=1)
add_bullet('Carga primero competencias, luego Gantt y por último participantes. Evita el orden inverso.')
add_bullet('Revisa la Carta Gantt después de cada carga para confirmar fechas.')
add_bullet('Al eliminar participantes, confirma dos veces: la acción también borra cuentas de usuario.')
add_bullet('Guarda backups manuales de las plantillas Excel que vas cargando.')
add_bullet('Antes de cerrar un programa, verifica que todos los informes estén generados y descargados.')
add_bullet('No compartas credenciales ni contraseñas visibles por canales inseguros.')

# ==========================
# 17. FAQ
# ==========================
add_heading('17. Preguntas frecuentes', level=1)

faqs = [
    ('Cargué los participantes pero no les llegó el correo',
     'Revisa la sección "Correos" para confirmar el estado del envío. Si el estado es "enviado" y el usuario no lo ve, pídele revisar la carpeta de spam. Si el estado es "fallido" o "parcial", la columna de error te indica el motivo (por ejemplo, correo inválido).'),
    ('¿Puedo editar un participante después de cargado?',
     'Sí. Desde la tabla de participantes puedes actualizar datos o reasignar colaboradores. Si hay que corregir muchos registros, es más rápido eliminar todo y volver a cargar el Excel.'),
    ('¿Cómo cambio la fecha de cierre de una encuesta?',
     'Abre el editor de la encuesta y modifica el campo "Fecha de cierre". También se actualiza automáticamente cuando cambias la Carta Gantt del programa.'),
    ('Un usuario reclama que no puede responder su encuesta',
     'Verifica que la encuesta esté activa y que la fecha de cierre no haya pasado. Si todo está correcto, pídele que limpie la caché del navegador o que use otro navegador.'),
    ('¿Qué pasa si elimino un cliente por error?',
     'Contacta a soporte técnico de inmediato. Los clientes tienen sus programas asociados, por lo que una eliminación puede arrastrar datos sensibles.'),
    ('¿Cuántos administradores puede haber?',
     'No hay un límite técnico, pero se recomienda mantener un grupo reducido y controlado dentro del equipo de MSO.'),
    ('¿La plataforma funciona desde celular para administrar?',
     'Las tareas básicas sí, pero para cargas masivas (Excel de participantes, Gantt, competencias) recomendamos siempre usar un computador.'),
]
for q, a in faqs:
    p = doc.add_paragraph()
    r = p.add_run(q)
    r.bold = True
    r.font.color.rgb = MORADO
    r.font.size = Pt(12)
    add_para(a, size=11)

# ==========================
# 18. SOPORTE
# ==========================
add_heading('18. Soporte y contacto', level=1)
add_para('Si tienes dudas o necesitas apoyo técnico, cuentas con estos canales:')
add_table_header(
    ['Canal', 'Detalle'],
    [
        ['Soporte técnico MSO', 'Problemas de acceso, errores de plataforma o sugerencias de mejora.'],
        ['Dashboard de Resend', 'Validación del envío de correos transaccionales.'],
        ['Módulo de Incidencias', 'Registro formal de cualquier problema detectado.'],
    ]
)

# ==========================
# CIERRE
# ==========================
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('MSO Chile · Modelos y Soluciones Organizacionales')
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = MORADO

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Plataforma de Transferencia al Puesto de Trabajo (TPT)')
r.font.size = Pt(10)
r.font.color.rgb = TEXTO_SEC

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('https://plataforma.msochile.cl  ·  Abril 2026')
r.font.size = Pt(10)
r.font.color.rgb = TEXTO_SEC

out = r"c:\Users\the_r\Documents\GitHub\plataforma-mso\docs\Guia_Administrador_TPT.docx"
doc.save(out)
print("OK:", out)
import os
print("Tamano:", round(os.path.getsize(out)/1024, 1), "KB")
