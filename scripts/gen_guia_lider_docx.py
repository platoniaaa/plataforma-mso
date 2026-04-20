"""Genera Guia_Lider_TPT.docx con estilos profesionales."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Colores MSO
MORADO = RGBColor(0x3D, 0x0C, 0x4B)
MORADO_CLARO = RGBColor(0x6B, 0x1D, 0x7B)
NARANJO = RGBColor(0xF5, 0x82, 0x20)
TEXTO = RGBColor(0x2D, 0x37, 0x48)
TEXTO_SEC = RGBColor(0x71, 0x80, 0x96)
BORDE = RGBColor(0xE2, 0xE8, 0xF0)

doc = Document()

# Configurar margenes
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Fuente por defecto
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = TEXTO

# ==========================
# HELPERS
# ==========================
def add_heading(text, level=1, color=None):
    if color is None:
        color = MORADO
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    if level == 1:
        run.font.size = Pt(22)
    elif level == 2:
        run.font.size = Pt(17)
    else:
        run.font.size = Pt(14)
    run.bold = True
    return h

def add_para(text, size=11, bold=False, color=None, align=None):
    if color is None:
        color = TEXTO
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    return p

def add_bullet(text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.font.color.rgb = TEXTO
    return p

def add_numbered(text, size=11):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.font.color.rgb = TEXTO
    return p

def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def add_callout(tipo, texto):
    """tipo: 'tip' | 'warning' | 'success'"""
    colores = {
        'tip':     ('E3F2FD', '1976D2', '💡 Consejo'),
        'warning': ('FFF3E0', 'F57C00', '⚠ Importante'),
        'success': ('E8F5E9', '388E3C', '✓ Tip')
    }
    bg, border_hex, titulo = colores[tipo]
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg)
    cell.paragraphs[0].clear()
    p_tit = cell.paragraphs[0]
    r_tit = p_tit.add_run(titulo + "\n")
    r_tit.bold = True
    r_tit.font.color.rgb = RGBColor.from_string(border_hex)
    r_tit.font.size = Pt(11)
    p_body = cell.add_paragraph()
    r_body = p_body.add_run(texto)
    r_body.font.size = Pt(10.5)
    r_body.font.color.rgb = TEXTO
    doc.add_paragraph()  # espacio

def add_step(numero, texto):
    """Agrega un paso numerado estilo tarjeta."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(1.2)
    table.columns[1].width = Cm(15.3)
    cell_num = table.cell(0, 0)
    cell_txt = table.cell(0, 1)
    set_cell_bg(cell_num, '3D0C4B')
    cell_num.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_num.paragraphs[0].clear()
    p_num = cell_num.paragraphs[0]
    p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_num = p_num.add_run(str(numero))
    r_num.bold = True
    r_num.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r_num.font.size = Pt(14)

    cell_txt.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_txt.paragraphs[0].clear()
    p_txt = cell_txt.paragraphs[0]
    r_txt = p_txt.add_run(texto)
    r_txt.font.size = Pt(11)
    r_txt.font.color.rgb = TEXTO
    doc.add_paragraph()

def add_table_header(headers, rows):
    """Agrega tabla con header morado y filas."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, '3D0C4B')
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(11)
    for row_idx, row_data in enumerate(rows, 1):
        row = table.rows[row_idx]
        for col_idx, val in enumerate(row_data):
            cell = row.cells[col_idx]
            if row_idx % 2 == 0:
                set_cell_bg(cell, 'FAFAFA')
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10.5)
            r.font.color.rgb = TEXTO
    doc.add_paragraph()

# ==========================
# PORTADA
# ==========================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MSO CHILE')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = MORADO

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Guía del Líder')
run.bold = True
run.font.size = Pt(40)
run.font.color.rgb = MORADO

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Plataforma de Transferencia al Puesto de Trabajo (TPT)')
run.font.size = Pt(16)
run.font.color.rgb = NARANJO
run.italic = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'Esta guía te acompañará en el uso de la plataforma durante todo el ciclo '
    'del programa de desarrollo de competencias. Encontrarás cómo acceder, '
    'responder evaluaciones, dar feedback a tu colaborador y seguir el avance '
    'del programa.'
)
run.font.size = Pt(12)
run.font.color.rgb = TEXTO_SEC
run.italic = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Versión 1.0  ·  Abril 2026')
run.font.size = Pt(11)
run.font.color.rgb = TEXTO_SEC

doc.add_page_break()

# ==========================
# TABLA DE CONTENIDOS
# ==========================
add_heading('Contenido', level=1)
toc_items = [
    '1. Primer acceso a la plataforma',
    '2. ¿Olvidaste tu contraseña?',
    '3. Tu pantalla principal: Mi Programa',
    '4. Cómo responder tus evaluaciones',
    '5. Tipos de evaluación en el programa',
    '6. Tu equipo y tu colaborador',
    '7. Cómo dar feedback a tu colaborador',
    '8. Archivos y recursos del programa',
    '9. Carta Gantt y cronograma',
    '10. Informes generados por la plataforma',
    '11. Reportar incidencias',
    '12. Preguntas frecuentes',
    '13. Soporte y contacto',
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(item)
    run.font.size = Pt(12)
    run.font.color.rgb = MORADO

doc.add_page_break()

# ==========================
# 1. ACCESO
# ==========================
add_heading('1. Primer acceso a la plataforma', level=1)
add_para(
    'Todos los líderes reciben un correo de bienvenida con sus credenciales cuando '
    'el administrador carga los datos del programa. Si todavía no lo recibiste, revisá '
    'tu carpeta de spam o contactá al administrador del programa.'
)

add_heading('Pasos para ingresar', level=2, color=MORADO_CLARO)
add_step(1, 'Abrí un navegador (Chrome, Edge, Firefox o Safari) e ingresá a: https://plataforma.msochile.cl')
add_step(2, 'Ingresá tu email corporativo en el campo correspondiente.')
add_step(3, 'Ingresá tu contraseña. Si es tu primer acceso, usá la contraseña que recibiste por correo.')
add_step(4, 'Click en el botón "Ingresar".')

add_callout('tip',
    'Te recomendamos cambiar tu contraseña en tu primer acceso por una que sea segura '
    'y fácil de recordar. Puedes hacerlo usando la opción "¿Olvidaste tu contraseña?" del login.')

# ==========================
# 2. RECUPERAR
# ==========================
add_heading('2. ¿Olvidaste tu contraseña?', level=1)
add_para('Si no recordás tu contraseña, seguí estos pasos:')

add_step(1, 'En la pantalla de login, click en "¿Olvidaste tu contraseña?"')
add_step(2, 'Ingresá tu email corporativo y click en "Enviar instrucciones".')
add_step(3, 'Recibirás un correo desde "MSO Plataforma <no_reply@msochile.cl>" con un link para restablecer tu contraseña.')
add_step(4, 'El link te lleva a una pantalla donde podés definir tu nueva contraseña. Elegí una combinación segura de al menos 6 caracteres.')

add_callout('warning',
    'El link de recuperación expira en 1 hora y solo puede usarse una vez por motivos '
    'de seguridad. Si no lo usás a tiempo, solicitá otro.')

# ==========================
# 3. MI PROGRAMA
# ==========================
add_heading('3. Tu pantalla principal: Mi Programa', level=1)
add_para(
    'Al iniciar sesión, tu pantalla principal es Mi Programa. Desde allí tienes una '
    'vista clara de todo el ciclo del programa en el que participas.'
)

add_heading('¿Qué encuentras en Mi Programa?', level=2, color=MORADO_CLARO)
add_bullet('Datos del programa: nombre, cliente, fechas de inicio y término.')
add_bullet('3 fases con pasos (steppers): Medición Inicial, Medición Final, Comparativo.')
add_bullet('Estado de cada paso: completada, en curso, pendiente o bloqueada.')
add_bullet('Botón "Responder" cuando una evaluación esté disponible.')

add_callout('tip',
    'Esta vista es tu centro de control. Si ves un paso en "Pendiente", significa que '
    'debes completarlo. Los pasos "Completada" ya fueron cerrados.')

# ==========================
# 4. EVALUACIONES
# ==========================
add_heading('4. Cómo responder tus evaluaciones', level=1)
add_para(
    'Las evaluaciones son el corazón del programa. Permiten medir las competencias y '
    'comportamientos antes (medición inicial) y después (medición final) de la intervención.'
)

add_heading('Responder una evaluación', level=2, color=MORADO_CLARO)
add_step(1, 'En el menú lateral, click en "Mis Evaluaciones".')
add_step(2, 'Verás dos secciones: "Encuestas Pendientes" y "Encuestas Completadas".')
add_step(3, 'En una encuesta pendiente, click en "Responder".')
add_step(4, 'Leé las instrucciones y respondé cada pregunta. La mayoría son de tipo escala de niveles (1 a 4).')

add_heading('Escala de niveles (1 a 4)', level=3, color=MORADO_CLARO)
add_bullet('Nivel 1: Conoce el concepto pero no lo aplica.')
add_bullet('Nivel 2: Aplica con orientación o apoyo.')
add_bullet('Nivel 3: Aplica de forma consistente y autónoma.')
add_bullet('Nivel 4: Es referente y modelo para otros.')

add_step(5, 'Una barra de progreso te indica cuántas preguntas llevás respondidas.')
add_step(6, 'Al terminar, click en "Enviar Respuestas". Recibirás un correo de confirmación.')

add_heading('¿Te equivocaste al responder?', level=2, color=MORADO_CLARO)
add_callout('success',
    'Si ya enviaste una evaluación y necesitas corregirla, ve a Mis Evaluaciones → '
    'Encuestas Completadas y click en "Rehacer evaluación". Tus respuestas anteriores '
    'serán reemplazadas por las nuevas.')

# ==========================
# 5. TIPOS DE EVAL
# ==========================
add_heading('5. Tipos de evaluación en el programa', level=1)
add_para(
    'Durante el programa vas a responder cuatro tipos de evaluaciones distribuidas en '
    'dos momentos: inicial y final.'
)

add_table_header(
    ['Tipo', 'Momento', 'Quién evalúa', 'A quién'],
    [
        ['Autoevaluación Inicial', 'Al comenzar', 'Líder (tú)', 'A sí mismo'],
        ['Coevaluación Inicial', 'Al comenzar', 'Colaborador asignado', 'A su líder (tú)'],
        ['Autoevaluación Final', 'Al cerrar', 'Líder (tú)', 'A sí mismo'],
        ['Coevaluación Final', 'Al cerrar', 'Colaborador asignado', 'A su líder (tú)'],
    ]
)

add_callout('tip',
    'La combinación entre tu autoevaluación y la coevaluación de tu colaborador '
    'permite identificar brechas de percepción y áreas de desarrollo específicas.')

# ==========================
# 6. MI EQUIPO
# ==========================
add_heading('6. Tu equipo y tu colaborador', level=1)
add_para(
    'Como líder, tenés un colaborador asignado que participará contigo durante el '
    'programa. Es tu "par de desarrollo" para los ejercicios de coevaluación y feedback.'
)

add_heading('Ver tu equipo', level=2, color=MORADO_CLARO)
add_step(1, 'En el menú lateral, click en "Mi Equipo".')
add_step(2, 'Verás una tarjeta con el nombre y cargo de tu colaborador, junto con un contador de feedback registrado.')

add_callout('warning',
    'Si no ves ningún colaborador asignado, contactá al administrador del programa.')

# ==========================
# 7. FEEDBACK
# ==========================
add_heading('7. Cómo dar feedback a tu colaborador', level=1)
add_para(
    'El feedback estructurado es una herramienta clave del programa. Te permite '
    'comunicar observaciones concretas a tu colaborador en tres dimensiones: '
    'fortalezas, aspectos a reforzar y recomendaciones.'
)

add_heading('Registrar un feedback', level=2, color=MORADO_CLARO)
add_step(1, 'Ve a "Mi Equipo".')
add_step(2, 'En la tarjeta de tu colaborador, click en "Dar Feedback".')
add_step(3, 'Completá los tres campos obligatorios: Fortaleza observada, Aspecto a reforzar y Recomendación.')
add_step(4, 'Click en "Enviar Feedback". Podrás ver el histórico de feedback entregado en la misma vista.')

add_heading('Recomendaciones para dar un buen feedback', level=2, color=MORADO_CLARO)
add_bullet('Sé específico: mencioná una situación concreta, no generalidades.')
add_bullet('Enfocá en el comportamiento observado, no en la persona.')
add_bullet('Ofrecé recomendaciones accionables.')
add_bullet('Mantené un tono constructivo y respetuoso.')

# ==========================
# 8. RECURSOS
# ==========================
add_heading('8. Archivos y recursos del programa', level=1)
add_para(
    'El administrador del programa puede subir documentos, lecturas, presentaciones '
    'y otros recursos para complementar tu aprendizaje.'
)
add_step(1, 'Ve a "Archivos y Recursos" en el menú lateral.')
add_step(2, 'Verás una lista de todos los archivos disponibles. Click en cada uno para descargarlo o visualizarlo.')

# ==========================
# 9. GANTT
# ==========================
add_heading('9. Carta Gantt y cronograma', level=1)
add_para(
    'La Carta Gantt del programa muestra todas las actividades planificadas, sus '
    'fechas y tu progreso a lo largo del tiempo.'
)
add_para('Incluye hitos como:')
add_bullet('Talleres y capacitaciones.')
add_bullet('Aplicación de autoevaluaciones y coevaluaciones.')
add_bullet('Checkpoints de acompañamiento grupal.')
add_bullet('Generación de informes individuales y consolidados.')
add_bullet('Cierre del programa.')
add_para(
    'La Carta Gantt te ayuda a saber qué viene a continuación y cuándo deberías '
    'completar cada actividad.'
)

# ==========================
# 10. INFORMES
# ==========================
add_heading('10. Informes generados por la plataforma', level=1)
add_para(
    'Una vez completada una medición (inicial o final), el administrador del programa '
    'genera informes con IA que resumen los resultados.'
)

add_heading('Tipos de informes', level=2, color=MORADO_CLARO)
add_table_header(
    ['Informe', 'Descripción'],
    [
        ['Informe Individual Inicial',
         'Resumen personal al comenzar el programa: niveles en cada competencia, brechas identificadas.'],
        ['Informe Consolidado Inicial',
         'Vista agregada del grupo de líderes en el programa.'],
        ['Informe Individual Final',
         'Resumen al cerrar el programa, con cambios respecto a la medición inicial.'],
        ['Informe Consolidado Final',
         'Vista agregada del grupo al cerrar el programa.'],
        ['Informe de Brechas',
         'Comparativo entre la medición inicial y final, mostrando avances concretos.'],
    ]
)

add_callout('tip',
    'Estos informes son generados por el administrador. Si querés acceder al tuyo, '
    'contactá al administrador del programa.')

# ==========================
# 11. INCIDENCIAS
# ==========================
add_heading('11. Reportar incidencias', level=1)
add_para(
    'Si encontrás algún error técnico, tenés una duda o querés sugerir una mejora, '
    'podés reportarlo directamente desde la plataforma.'
)
add_step(1, 'En el menú lateral, click en "Incidencias".')
add_step(2, 'Completá los campos: categoría, título, descripción del problema y, opcionalmente, adjuntá una captura de pantalla.')
add_step(3, 'Click en "Reportar". El equipo de soporte recibirá tu reporte y te contactará.')

# ==========================
# 12. FAQ
# ==========================
add_heading('12. Preguntas frecuentes', level=1)

faqs = [
    ('No me llegó el correo de bienvenida, ¿qué hago?',
     'Revisá tu carpeta de spam o correo no deseado. Si tampoco está ahí, contactá al administrador del programa para que verifique el envío.'),
    ('El link de reset de contraseña dice que expiró',
     'Los links de reset son válidos por 1 hora y solo se pueden usar una vez. Si expiró, solicitá otro desde "¿Olvidaste tu contraseña?" del login.'),
    ('No veo mi evaluación disponible',
     'Las evaluaciones se activan según el cronograma del programa. Si la fecha ya pasó y no la ves, contactá al administrador para que verifique el estado de activación.'),
    ('Respondí una evaluación por error, ¿puedo corregirla?',
     'Sí. Ve a Mis Evaluaciones → Encuestas Completadas y click en "Rehacer evaluación". Tus respuestas anteriores serán reemplazadas.'),
    ('No veo a mi colaborador asignado',
     'Contactá al administrador del programa. Él puede verificar la asignación y corregirla si es necesario.'),
    ('¿Puedo usar la plataforma desde el celular?',
     'Sí, la plataforma funciona en celulares y tablets. Para completar evaluaciones largas recomendamos usar computador o tablet para mayor comodidad.'),
    ('¿Mis respuestas son confidenciales?',
     'Tus autoevaluaciones y las coevaluaciones de tu colaborador se usan exclusivamente para los informes del programa. El tratamiento es agregado en los informes consolidados y el individual solo lo ven el administrador del programa y tú.'),
]
for q, a in faqs:
    p = doc.add_paragraph()
    r = p.add_run(q)
    r.bold = True
    r.font.color.rgb = MORADO
    r.font.size = Pt(12)
    add_para(a, size=11)

# ==========================
# 13. SOPORTE
# ==========================
add_heading('13. Soporte y contacto', level=1)
add_para('Si necesitás ayuda o tenés alguna duda que no aparece en esta guía, contactá a:')
add_table_header(
    ['Canal', 'Detalle'],
    [
        ['Administrador del programa', 'Consultas sobre el contenido, fechas o activación de evaluaciones.'],
        ['Soporte técnico MSO', 'Problemas de acceso, errores técnicos o sugerencias sobre la plataforma.'],
        ['Módulo de Incidencias', 'Reportá dentro de la plataforma cualquier error técnico.'],
    ]
)

# ==========================
# FOOTER / CIERRE
# ==========================
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('MSO Chile · Modelos y Soluciones Organizacionales')
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = MORADO

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Plataforma de Transferencia al Puesto de Trabajo (TPT)')
r.font.size = Pt(10)
r.font.color.rgb = TEXTO_SEC

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('https://plataforma.msochile.cl  ·  Abril 2026')
r.font.size = Pt(10)
r.font.color.rgb = TEXTO_SEC

# Guardar
out = r"c:\Users\the_r\Documents\GitHub\plataforma-mso\docs\Guia_Lider_TPT.docx"
doc.save(out)
print("OK:", out)
import os
print("Tamano:", round(os.path.getsize(out)/1024, 1), "KB")
